from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGIN_DXA = {"top": 80, "bottom": 80, "start": 120, "end": 120}
ACCENT_BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F2F4F7"
INK = "0B2545"
MUTED = "667085"
CODE_FILL = "F6F8FA"


def set_font(run, name="Calibri", size=11, color="000000", bold=None, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_bidi(element, value=True):
    props = element.get_or_add_pPr() if hasattr(element, "get_or_add_pPr") else element
    tag = qn("w:bidi")
    node = props.find(tag)
    if value and node is None:
        props.append(OxmlElement("w:bidi"))
    elif not value and node is not None:
        props.remove(node)


def set_paragraph(paragraph, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph.alignment = align
    paragraph.paragraph_format.widow_control = True
    if rtl:
        set_bidi(paragraph._p)


def set_cell_margins(cell, margins=CELL_MARGIN_DXA):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in margins.items():
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(TABLE_INDENT_DXA))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.first_child_found_in("w:shd")
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="D0D5DD", size="6"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)
    set_font(run, size=9, color=MUTED)


def configure_document(document):
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string("000000")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    heading_tokens = {
        "Heading 1": (16, ACCENT_BLUE, 16, 8),
        "Heading 2": (13, ACCENT_BLUE, 12, 6),
        "Heading 3": (12, DARK_BLUE, 8, 4),
    }
    for style_name, (size, color, before, after) in heading_tokens.items():
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.widow_control = True

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:cs"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    header_p = header.paragraphs[0]
    header_p.text = "PREMYA  |  מסמך אפיון מערכת"
    set_paragraph(header_p, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    for run in header_p.runs:
        set_font(run, size=9, color=MUTED, bold=True)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.text = "מסמך אפיון – מערכת לניהול שיטות פרמיה  |  עמוד "
    set_paragraph(footer_p, rtl=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for run in footer_p.runs:
        set_font(run, size=9, color=MUTED)
    add_page_number(footer_p)


def add_inline_runs(paragraph, text, code=False):
    if code:
        run = paragraph.add_run(text)
        set_font(run, name="Consolas", size=9, color=INK)
        return
    pattern = re.compile(r"(`[^`]+`|\*\*[^*]+\*\*|\*[^*]+\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            run = paragraph.add_run(text[position:match.start()])
            set_font(run)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_font(run, name="Consolas", size=9, color=INK)
        else:
            run = paragraph.add_run(token.strip("*"))
            set_font(run, bold=token.startswith("**"))
        position = match.end()
    if position < len(text):
        run = paragraph.add_run(text[position:])
        set_font(run)


def add_heading(document, text, level):
    normalized_level = min(max(level, 1), 3)
    paragraph = document.add_paragraph(style=f"Heading {normalized_level}")
    set_paragraph(paragraph, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_inline_runs(paragraph, text)
    return paragraph


def add_body_paragraph(document, text, align=WD_ALIGN_PARAGRAPH.RIGHT):
    paragraph = document.add_paragraph()
    set_paragraph(paragraph, rtl=True, align=align)
    add_inline_runs(paragraph, text)
    return paragraph


def add_list_item(document, text, numbered=False):
    style = "List Number" if numbered else "List Bullet"
    paragraph = document.add_paragraph(style=style)
    set_paragraph(paragraph, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    add_inline_runs(paragraph, text)
    return paragraph


def add_code_block(document, lines):
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.15)
        paragraph.paragraph_format.right_indent = Inches(0.15)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing = 1.0
        set_paragraph(paragraph, rtl=False, align=WD_ALIGN_PARAGRAPH.LEFT)
        p_pr = paragraph._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), CODE_FILL)
        p_pr.append(shd)
        run = paragraph.add_run(line if line else " ")
        set_font(run, name="Consolas", size=8.5, color=INK)


def clean_cell_text(value):
    return re.sub(r"\s+", " ", value.strip().strip("`"))


def is_table_separator(line):
    cells = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return bool(cells) and all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells)


def parse_table(lines, start):
    header = [clean_cell_text(x) for x in lines[start].strip().strip("|").split("|")]
    separator = start + 1
    if separator >= len(lines) or not is_table_separator(lines[separator]):
        return None, start
    rows = []
    index = separator + 1
    while index < len(lines) and "|" in lines[index] and lines[index].strip():
        rows.append([clean_cell_text(x) for x in lines[index].strip().strip("|").split("|")])
        index += 1
    column_count = len(header)
    header = (header + [""] * column_count)[:column_count]
    rows = [(row + [""] * column_count)[:column_count] for row in rows]
    return (header, rows), index


def add_markdown_table(document, header, rows):
    column_count = len(header)
    table = document.add_table(rows=1, cols=column_count)
    widths = [CONTENT_WIDTH_DXA // column_count] * column_count
    widths[-1] += CONTENT_WIDTH_DXA - sum(widths)
    set_table_geometry(table, widths)
    set_table_borders(table)
    repeat_table_header(table.rows[0])
    for index, value in enumerate(header):
        cell = table.rows[0].cells[index]
        shade_cell(cell, LIGHT_BLUE)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        set_paragraph(paragraph, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        run = paragraph.add_run(value)
        set_font(run, size=9.5, color=INK, bold=True)
    for row_values in rows:
        cells = table.add_row().cells
        for index, value in enumerate(row_values):
            cells[index].text = ""
            paragraph = cells[index].paragraphs[0]
            set_paragraph(paragraph, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
            add_inline_runs(paragraph, value)
            for run in paragraph.runs:
                run.font.size = Pt(9)
    document.add_paragraph().paragraph_format.space_after = Pt(2)


def add_toc(document, headings):
    heading = add_heading(document, "תוכן עניינים", 1)
    heading.paragraph_format.page_break_before = False
    for number, title in headings:
        paragraph = document.add_paragraph(style="List Number")
        set_paragraph(paragraph, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
        add_inline_runs(paragraph, f"{number}  {title}")
    document.add_page_break()


def convert_markdown(markdown_path, output_path):
    source = Path(markdown_path).read_text(encoding="utf-8")
    lines = source.splitlines()
    document = Document()
    configure_document(document)

    title = "מסמך אפיון – מערכת לניהול שיטות פרמיה"
    add_title = document.add_paragraph()
    add_title.paragraph_format.space_before = Pt(28)
    add_title.paragraph_format.space_after = Pt(8)
    set_paragraph(add_title, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    title_run = add_title.add_run(title)
    set_font(title_run, size=23, color=INK, bold=True)

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(20)
    set_paragraph(subtitle, rtl=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    subtitle_run = subtitle.add_run("מסמך דרישות, מבנה נתונים, ארכיטקטורה ותכנון API  |  גרסה 1.0")
    set_font(subtitle_run, size=12, color=MUTED)

    headings = []
    for line in lines:
        match = re.match(r"^##\s+(\d+)\.\s+(.+)$", line.strip())
        if match:
            headings.append((match.group(1) + ".", match.group(2).strip()))
    add_toc(document, headings)

    index = 0
    in_code = False
    code_lines = []
    while index < len(lines):
        raw = lines[index]
        stripped = raw.strip()
        if index == 0 and stripped.startswith("# "):
            index += 1
            continue
        if stripped == "## תוכן עניינים":
            index += 1
            while index < len(lines) and not lines[index].startswith("## 1."):
                index += 1
            continue
        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            index += 1
            continue
        if in_code:
            code_lines.append(raw)
            index += 1
            continue
        if not stripped:
            index += 1
            continue
        heading_match = re.match(r"^(#{2,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1)) - 1
            add_heading(document, heading_match.group(2).strip(), level)
            index += 1
            continue
        if stripped.startswith("|") and index + 1 < len(lines):
            parsed, next_index = parse_table(lines, index)
            if parsed:
                add_markdown_table(document, *parsed)
                index = next_index
                continue
        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered_match:
            add_list_item(document, numbered_match.group(1), numbered=True)
            index += 1
            continue
        bullet_match = re.match(r"^[-*]\s+(.+)$", stripped)
        if bullet_match:
            add_list_item(document, bullet_match.group(1), numbered=False)
            index += 1
            continue
        if stripped.startswith("> "):
            paragraph = add_body_paragraph(document, stripped[2:])
            paragraph.paragraph_format.left_indent = Inches(0.25)
            paragraph.paragraph_format.right_indent = Inches(0.25)
            for run in paragraph.runs:
                run.italic = True
            index += 1
            continue
        add_body_paragraph(document, stripped)
        index += 1
    if in_code and code_lines:
        add_code_block(document, code_lines)

    properties = document.core_properties
    properties.title = title
    properties.subject = "מסמך אפיון טכני"
    properties.author = "Premya"
    properties.keywords = "Premya, אפיון, פרמיה, Excel, API"
    document.save(output_path)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_spec_docx.py <input.md> <output.docx>")
    convert_markdown(sys.argv[1], sys.argv[2])
